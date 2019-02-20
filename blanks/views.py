# -*- coding: utf-8 -*-
from __future__ import unicode_literals
from django.views.generic import TemplateView
from django.shortcuts import render, redirect, get_object_or_404
from django.core.files.storage import FileSystemStorage
from django.shortcuts import redirect
from django.core.files import File

from .forms import DocumentForm, RawProductionForm, VariablesForm, EditedDocumentForm
from .models import DocFile, DocFields

import os
import re
from docx import Document

from django.contrib import messages


def Home(request):
    return render(request, 'home.html')


def file_list(request):
    documents = DocFile.objects.all()
    return render(request, 'file_list.html',{
        'documents': documents
    })


def upload_files(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect("file_list")
    else:
        form = DocumentForm()
    return render(request, 'upload_files.html',{
        'form': form
    })


def docx_words_replace(doc_obj, regex, replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs


            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_words_replace(cell, regex, replace)

def edit_files(request, file_id):
    instance = get_object_or_404(DocFile, id=file_id)
    exact_file = Document(instance.document)
    variables = []
    for paragraph in exact_file.paragraphs:
        match = re.findall(r"\{(.*?)\}", paragraph.text)
        variables.append(match)
    for table in exact_file.tables:
        for row in table.rows:
            for cell in row.cells:
                match = re.findall(r"\{(.*?)\}", cell.text)
                variables.append(match)
    exact_file.save('green.pdf')
    temporary_list =  []
    for i in variables:
        for j in i :
            if len(j) > 0:
                temporary_list.append(j)
    variables = temporary_list
    #  https://stackoverflow.com/questions/9792664/set-changes-element-order
    variables_set = sorted(set(variables), key=variables.index)

    inputs_amount = len(variables)
    print(variables)
    inputs_list = []
    my_form = VariablesForm(request.POST, variables=variables_set)
    # Start changing variables
    if request.method == 'POST' and my_form.is_valid():
        print(my_form)
        input_texts = my_form.get_input_text()
        for i, input_text in input_texts:
           inputs_list.append(input_text)

    print(inputs_list)
    my_dict = dict(zip(variables_set, inputs_list))
    print(my_dict)
    for word , replacement in my_dict.items():
        word_re = re.compile(word)
        docx_words_replace(exact_file, word_re, replacement)
    regex1 = re.compile(r"{")
    regex2 = re.compile(r"}")
    replace = r""
    docx_words_replace(exact_file, regex1, replace)
    docx_words_replace(exact_file, regex2, replace)

    if len(inputs_list) != 0:
        contract_name = inputs_list[0]
        files_full_name = '{}.docx'.format(contract_name)
        desktop = os.path.normpath(os.path.expanduser("~/Downloads/aca_docs"))
        print(desktop,'kkkkkkkkkkkkkk' )
        try:
            os.makedirs(desktop)
        except FileExistsError:
            # directory already exists
            pass
        exact_file.save(os.path.join(desktop, files_full_name))
        messages.success(request, 'Saved in {}'.format(desktop))
        #  exact_file.save(os.path.join(desktop, files_full_name))
        #  html_name = 'Saved in <a href="file:///home/grigor/Downloads/Aca_docs/{}" target="_blank">downloads</a>'.format(
        #     files_full_name)
        # messages.success(request, html_name)
        # nm = EditedDocumentForm(files_full_name, File(exact_file))
        # nm.save()
    return render(request, 'edit_files.html', context={'variables': variables, "form": my_form})


def delete_book(request, pk):
    if request.method == 'POST':
        got_to_delete = DocFile.objects.get(pk=pk)
        got_to_delete.delete()
    return redirect('file_list')

